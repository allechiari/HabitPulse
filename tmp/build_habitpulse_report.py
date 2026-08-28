from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\a.chiari\source\repos\Habit")
OUT = ROOT / "docs" / "relazione"
OUT.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT / "Relazione_HabitPulse.docx"

BLUE = "2E74B5"
DARK = "1F4D78"
LIGHT = "F2F4F7"
MUTED = "666666"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, DARK),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(5)
    bullet.paragraph_format.line_spacing = 1.10

    number = doc.styles["List Number"]
    number.font.name = "Calibri"
    number.font.size = Pt(11)
    number.paragraph_format.left_indent = Inches(0.5)
    number.paragraph_format.first_line_indent = Inches(-0.25)
    number.paragraph_format.space_after = Pt(5)
    number.paragraph_format.line_spacing = 1.10

    footer = section.footer.paragraphs[0]
    footer.style = normal
    footer.paragraph_format.space_after = Pt(0)
    add_page_number(footer)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, label, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(label + ": ")
    r.bold = True
    p.add_run(text)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.right_indent = Inches(0.35)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(9.5)


def add_index_line(doc, number, title, page, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.25))
    run = p.add_run(f"{number}   {title}")
    run.bold = indent == 0
    p.add_run(f"\t{page}")


doc = Document()
style_document(doc)

# Pagina 1 - titolo e indice
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(38)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("HabitPulse")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor.from_string(DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Relazione di progetto")
r.font.size = Pt(15)
r.font.color.rgb = RGBColor.from_string(MUTED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(30)
r = p.add_run("Alessandro Chiari")
r.font.size = Pt(12)

add_heading(doc, "Indice", 1)
for row in [
    ("1", "Introduzione", "2", 0),
    ("2", "Obiettivo", "2", 0),
    ("3", "Funzionalità principali", "2", 0),
    ("3.1", "Accesso e profilo", "2", 0.28),
    ("3.2", "Dashboard e attività", "2", 0.28),
    ("3.3", "Progressi e analisi", "2", 0.28),
    ("4", "Architettura e database", "3", 0),
    ("5", "Tecnologie utilizzate", "4", 0),
    ("6", "Sicurezza e stato del progetto", "4", 0),
    ("7", "Verifiche effettuate", "5", 0),
    ("8", "Manuale d'uso e installazione", "6", 0),
]:
    add_index_line(doc, *row)

doc.add_page_break()

# Pagina 2
add_heading(doc, "1  Introduzione", 1)
doc.add_paragraph(
    "HabitPulse è un'applicazione web full stack dedicata alla gestione delle abitudini quotidiane. "
    "Permette a ogni utente di registrare attività positive, chiamate habit, e comportamenti da limitare, "
    "chiamati vizi. Il progetto utilizza un'architettura client-server basata su React, Node.js, Express e MongoDB."
)
add_heading(doc, "2  Obiettivo", 1)
doc.add_paragraph("L'obiettivo è offrire uno strumento semplice per:")
for text in [
    "creare e organizzare le proprie attività quotidiane;",
    "registrare un contatore e un obiettivo giornaliero;",
    "distinguere le buone abitudini dai comportamenti da ridurre;",
    "osservare l'andamento nel tempo attraverso riepiloghi e grafici;",
    "mantenere separati i dati dei diversi utenti.",
]:
    p = doc.add_paragraph(text, style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

add_heading(doc, "3  Funzionalità principali", 1)
add_heading(doc, "3.1  Accesso e profilo", 2)
add_bullet(doc, "Registrazione", "creazione di un account con nome, cognome, email, password e telefono opzionale.")
add_bullet(doc, "Login", "accesso tramite email e password; la password viene confrontata con l'hash salvato nel database.")
add_bullet(doc, "Profilo", "modifica dei dati personali, cambio della password, logout e cancellazione dell'account.")

add_heading(doc, "3.2  Dashboard e attività", 2)
add_bullet(doc, "Creazione", "inserimento di titolo, tipo, descrizione, date, colore, unità di misura e obiettivo.")
add_bullet(doc, "Gestione", "modifica del colore e dell'obiettivo, riordinamento tramite drag and drop e arresto di un'attività.")
add_bullet(doc, "Visualizzazione", "possibilità di mostrare o nascondere le attività terminate.")

add_heading(doc, "3.3  Progressi e analisi", 2)
add_bullet(doc, "Progresso giornaliero", "aggiornamento del contatore e del target della giornata.")
add_bullet(doc, "Analisi", "filtri settimanali, mensili, annuali o personalizzati, grafici percentuali, copertura dei dati e calendario riepilogativo.")

doc.add_page_break()

# Pagina 3
add_heading(doc, "4  Architettura e database", 1)
doc.add_paragraph(
    "Il frontend comunica con il backend tramite API REST che scambiano dati in formato JSON. "
    "Il server applica la logica dell'applicazione e usa Mongoose per leggere e scrivere i documenti su MongoDB Atlas."
)

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.style = "Table Grid"
headers = table.rows[0].cells
headers[0].text = "Componente"
headers[1].text = "Responsabilità"
for cell in headers:
    set_cell_shading(cell, LIGHT)
    cell.paragraphs[0].runs[0].bold = True
for left, right in [
    ("Client React", "Interfaccia, navigazione, form, dashboard e grafici."),
    ("API Express", "Rotte per autenticazione, habit, progressi e analisi."),
    ("MongoDB", "Persistenza di utenti, attività e valori giornalieri."),
]:
    cells = table.add_row().cells
    cells[0].text = left
    cells[1].text = right
    for cell in cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
set_table_geometry(table, [2700, 6660])

add_heading(doc, "4.1  Users", 2)
doc.add_paragraph(
    "La collezione degli utenti contiene nome, cognome, email univoca, password cifrata, telefono e date di creazione e aggiornamento."
)
add_heading(doc, "4.2  Habits", 2)
doc.add_paragraph(
    "Ogni attività è collegata a un utente e contiene titolo, tipo (habit o vice), descrizione, data iniziale e finale, "
    "colore, ordine, obiettivo predefinito, unità di misura e stato di arresto."
)
add_heading(doc, "4.3  DailyProgress", 2)
doc.add_paragraph(
    "Il progresso giornaliero collega un'attività a una data e salva contatore e target. "
    "Un indice composto impedisce di avere due record per la stessa attività nello stesso giorno."
)

add_heading(doc, "4.4  Principali endpoint", 2)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.LEFT
for idx, text in enumerate(("Area", "Metodo", "Endpoint")):
    table.rows[0].cells[idx].text = text
    set_cell_shading(table.rows[0].cells[idx], LIGHT)
    table.rows[0].cells[idx].paragraphs[0].runs[0].bold = True
for row in [
    ("Autenticazione", "POST", "/api/auth/register - /login"),
    ("Attività", "GET/POST/PATCH", "/api/habits"),
    ("Progressi", "GET/PATCH", "/api/progress"),
    ("Analisi", "GET", "/api/analysis"),
]:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
set_table_geometry(table, [2500, 2200, 4660])

doc.add_page_break()

# Pagina 4
add_heading(doc, "5  Tecnologie utilizzate", 1)
add_bullet(doc, "Frontend", "React 19, Vite, React Router DOM e Recharts.")
add_bullet(doc, "Backend", "Node.js, Express.js, bcryptjs, CORS e dotenv.")
add_bullet(doc, "Database", "MongoDB Atlas con ODM Mongoose.")
add_bullet(doc, "Sviluppo", "ESLint, Nodemon e Concurrently.")

add_heading(doc, "6  Sicurezza e stato del progetto", 1)
doc.add_paragraph(
    "Le password vengono cifrate con bcrypt prima del salvataggio. In fase di registrazione viene richiesta una password "
    "di almeno otto caratteri, con maiuscola, minuscola, numero e carattere speciale. L'email è unica nel database."
)
doc.add_paragraph(
    "Il progetto è funzionante come prototipo locale, ma alcune parti devono essere completate prima di una pubblicazione reale:"
)
for text in [
    "introdurre autenticazione e autorizzazione server-side tramite token o sessione;",
    "ricavare l'identità dell'utente dalla sessione invece di accettare liberamente userId dal client;",
    "spostare l'indirizzo delle API in una variabile di ambiente;",
    "limitare CORS alle origini autorizzate e aggiungere protezione dai tentativi ripetuti di login;",
    "eliminare in modo coerente habit e progressi quando viene cancellato un account;",
    "rimuovere e ruotare qualsiasi credenziale di database comparsa nel codice o nella cronologia Git;",
    "definire in modo esplicito il fuso orario usato per i progressi giornalieri.",
]:
    p = doc.add_paragraph(text, style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

add_heading(doc, "6.1  Gestione dei progressi", 2)
doc.add_paragraph(
    "Attualmente l'apertura della dashboard crea automaticamente il progresso del giorno con valore zero. "
    "Prima di considerare definitivo il calcolo della copertura, è necessario decidere se un giorno debba risultare "
    "tracciato solo dopo un'azione esplicita dell'utente."
)

doc.add_page_break()

# Pagina 5
add_heading(doc, "7  Verifiche effettuate", 1)
doc.add_paragraph(
    "Durante la revisione del progetto sono stati eseguiti controlli automatici sul frontend e sul backend. "
    "La tabella seguente riporta esclusivamente verifiche realmente eseguite sul codice attuale."
)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.LEFT
for idx, text in enumerate(("Controllo", "Risultato", "Nota")):
    table.rows[0].cells[idx].text = text
    set_cell_shading(table.rows[0].cells[idx], LIGHT)
    table.rows[0].cells[idx].paragraphs[0].runs[0].bold = True
for row in [
    ("ESLint frontend", "SUPERATO", "Nessun errore segnalato."),
    ("Build Vite", "SUPERATO", "Build di produzione completata."),
    ("Sintassi Node.js", "SUPERATO", "Controller, modelli, rotte e configurazione validi."),
    ("Test automatici", "ASSENTI", "Non è presente una suite di test nel repository."),
]:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
set_table_geometry(table, [2600, 1900, 4860])

add_heading(doc, "7.1  Test consigliati", 2)
for text in [
    "registrazione, login, modifica profilo e cambio password;",
    "isolamento dei dati tra due utenti diversi;",
    "creazione, modifica, arresto e riordinamento delle attività;",
    "aggiornamento concorrente dei progressi giornalieri;",
    "calcolo delle percentuali per habit e vizi;",
    "intervalli di date, cambio del giorno e fusi orari;",
    "cancellazione completa dell'account e dei dati collegati.",
]:
    p = doc.add_paragraph(text, style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

add_heading(doc, "7.2  Nota sulle prestazioni", 2)
doc.add_paragraph(
    "La build segnala un bundle JavaScript principale superiore a 500 kB. Il comportamento non impedisce l'avvio, "
    "ma in futuro è consigliabile caricare in modo differito la pagina di analisi e la libreria dei grafici."
)

doc.add_page_break()

# Pagina 6
add_heading(doc, "8  Manuale d'uso e installazione", 1)
doc.add_paragraph("Per eseguire HabitPulse in locale sono necessari Node.js, npm e un database MongoDB raggiungibile.")

add_heading(doc, "8.1  Installazione", 2)
doc.add_paragraph("Dalla cartella principale del progetto installare le dipendenze dei tre livelli:")
add_code(doc, "npm install\ncd client && npm install\ncd ../server && npm install")

add_heading(doc, "8.2  Configurazione del server", 2)
doc.add_paragraph("Creare un file .env nella cartella server senza inserirlo nel repository:")
add_code(doc, "PORT=5000\nMONGO_URI=mongodb+srv://<utente>:<password>@<cluster>/habitpulse")

add_heading(doc, "8.3  Avvio", 2)
doc.add_paragraph("Per avviare frontend e backend insieme, dalla cartella principale eseguire:")
add_code(doc, "npm run dev")
doc.add_paragraph(
    "Il frontend è normalmente disponibile su http://localhost:5173, mentre il backend risponde su "
    "http://localhost:5000."
)

add_heading(doc, "8.4  Utilizzo essenziale", 2)
for text in [
    "Registrare un nuovo account oppure effettuare il login.",
    "Creare una nuova attività dalla dashboard, scegliendo se si tratta di un habit o di un vizio.",
    "Aggiornare ogni giorno contatore e obiettivo.",
    "Usare il trascinamento per modificare l'ordine delle schede.",
    "Aprire la pagina Analysis per consultare grafici e riepiloghi del periodo selezionato.",
    "Aprire Profile per aggiornare i dati, cambiare password, uscire o cancellare l'account.",
]:
    p = doc.add_paragraph(text, style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.core_properties.title = "Relazione HabitPulse"
doc.core_properties.subject = "Relazione semplice del progetto HabitPulse"
doc.core_properties.author = "Alessandro Chiari"
doc.core_properties.keywords = "HabitPulse, React, Express, MongoDB, relazione"
doc.save(DOCX_PATH)
print(DOCX_PATH)
